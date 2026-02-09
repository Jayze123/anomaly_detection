from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt


def _write_main_doc(path: Path) -> None:
    doc = Document()
    h = doc.add_heading("Anomaly Detection Application Documentation", level=0)
    h.runs[0].font.size = Pt(22)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "This application implements a modular anomaly inspection pipeline with a web UI. "
        "It supports category-specific anomaly detection on MVTec AD-style data and offers "
        "a user camera interface and an admin upload interface."
    )

    doc.add_heading("2. Project Structure", level=1)
    doc.add_paragraph("Key folders and files:")
    doc.add_paragraph("- configs/base.yaml: Primary configuration file.")
    doc.add_paragraph("- src/: Application source code.")
    doc.add_paragraph("  - src/api.py: FastAPI server with user/admin UIs.")
    doc.add_paragraph("  - src/pipeline.py: Batch pipeline.")
    doc.add_paragraph("  - src/models/mean_diff.py: Baseline anomaly model.")
    doc.add_paragraph("  - src/data/mvtec.py: MVTec data loader.")
    doc.add_paragraph("  - src/postproc/: Heatmap, mask, and bbox utilities.")
    doc.add_paragraph("  - src/vlm/semantics.py: LLaVA integration.")
    doc.add_paragraph("  - src/risk/: RPM lookup + policy mapping.")
    doc.add_paragraph("  - src/uncertainty/: confidence and review rules.")
    doc.add_paragraph("- artifacts/: Generated heatmaps, masks, and uploads.")
    doc.add_paragraph("- data/: Dataset folder.")
    doc.add_paragraph("- requirements.txt: Python dependencies.")

    doc.add_heading("3. Configuration", level=1)
    doc.add_paragraph("Important sections in configs/base.yaml:")
    doc.add_paragraph("- data: dataset root and category selection.")
    doc.add_paragraph("- paths: artifacts and outputs.")
    doc.add_paragraph("- postproc: heatmap normalization and thresholding.")
    doc.add_paragraph("- labels: category-specific label sets.")
    doc.add_paragraph("- risk: RPM table and risk-to-action mapping.")
    doc.add_paragraph("- uncertainty: confidence fusion and review threshold.")

    doc.add_heading("4. Web UI and Endpoints", level=1)
    doc.add_paragraph("The FastAPI app exposes:")
    doc.add_paragraph("- User UI: GET / (camera interface).")
    doc.add_paragraph("- Admin UI: GET /admin (upload interface).")
    doc.add_paragraph("- Upload endpoint: POST /upload.")
    doc.add_paragraph("- Analyze endpoint: POST /analyze.")

    doc.add_heading("5. Baseline Anomaly Detection", level=1)
    doc.add_paragraph(
        "Current baseline uses a mean-difference model: per-pixel mean and std are computed "
        "from training images and anomalies are scored by normalized deviation."
    )

    doc.add_heading("6. VLM, Risk, and Uncertainty", level=1)
    doc.add_paragraph(
        "VLM semantics are implemented with LLaVA-1.6 (Mistral). "
        "Risk uses a deterministic RPM lookup, and uncertainty combines anomaly and VLM confidence."
    )

    doc.add_heading("7. Running the Application", level=1)
    doc.add_paragraph("Install dependencies:")
    doc.add_paragraph("python -m pip install -r requirements.txt")
    doc.add_paragraph("Start the server:")
    doc.add_paragraph("python -m uvicorn src.api:app --reload")

    doc.save(str(path))


def _write_user_guide(path: Path) -> None:
    doc = Document()
    h = doc.add_heading("User Guide: Anomaly Inspection Application", level=0)
    h.runs[0].font.size = Pt(22)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "The system provides a user camera UI and an admin upload UI to run anomaly inspection."
    )

    doc.add_heading("2. Module Guide", level=1)
    doc.add_heading("2.1 src/api.py — Web API + UI", level=2)
    doc.add_paragraph("Routes: GET /, GET /admin, POST /upload, POST /analyze.")
    doc.add_heading("2.2 src/pipeline.py — Batch Pipeline", level=2)
    doc.add_paragraph("Batch processing for dataset evaluation.")
    doc.add_heading("2.3 src/data/mvtec.py — Dataset Loader", level=2)
    doc.add_paragraph("Loads MVTec AD style data and ground-truth masks.")
    doc.add_heading("2.4 src/models/mean_diff.py — Baseline Model", level=2)
    doc.add_paragraph("Mean/std baseline model for anomaly scoring.")
    doc.add_heading("2.5 src/postproc/* — Heatmap, Mask, BBoxes", level=2)
    doc.add_paragraph("Normalize, threshold, and bounding box extraction.")
    doc.add_heading("2.6 src/vlm/semantics.py — VLM", level=2)
    doc.add_paragraph("LLaVA-1.6 (Mistral) integration for defect labels.")
    doc.add_heading("2.7 src/risk/* — Risk Lookup + Policy", level=2)
    doc.add_paragraph("RPM lookup and risk-to-action mapping.")
    doc.add_heading("2.8 src/uncertainty/* — Confidence + Review", level=2)
    doc.add_paragraph("Confidence fusion and human review rules.")
    doc.add_heading("2.9 configs/base.yaml — Configuration", level=2)
    doc.add_paragraph("Central configuration for categories and thresholds.")

    doc.add_heading("3. End-to-End Process", level=1)
    doc.add_paragraph("User UI flow: start camera, capture, analyze, review result.")
    doc.add_paragraph("Admin UI flow: upload image, analyze, review result.")

    doc.add_heading("4. Outputs", level=1)
    doc.add_paragraph("Heatmaps, masks, and JSON outputs saved under artifacts/ and outputs/.")

    doc.add_heading("5. Running", level=1)
    doc.add_paragraph("Install dependencies and run: python -m uvicorn src.api:app --reload")

    doc.save(str(path))


def main() -> int:
    docs_dir = Path("docs")
    docs_dir.mkdir(parents=True, exist_ok=True)
    _write_main_doc(docs_dir / "Anomaly_Detection_Documentation.docx")
    _write_user_guide(docs_dir / "User_Guide.docx")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
