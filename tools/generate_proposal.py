from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt


def main() -> int:
    out_path = Path("docs/Dissertation_Proposal.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    h = doc.add_heading(
        "MSc Robotics Dissertation Proposal: Interpretable Industrial Image Anomaly Inspection",
        level=0,
    )
    h.runs[0].font.size = Pt(20)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This proposal outlines the design and implementation of an industrial image anomaly "
        "inspection pipeline with interpretable semantics and safety-aligned risk scoring. "
        "The system targets MVTec AD categories (bottle, cable, wood, tile, leather) and outputs "
        "anomaly detection, localization, defect labels, evidence, risk class, and action "
        "recommendations. The work combines state-of-the-art anomaly detection baselines with "
        "a vision-language model for defect semantics and deterministic risk rules."
    )

    doc.add_heading("1. Background and Motivation", level=1)
    doc.add_paragraph(
        "Industrial visual inspection often relies on scarce defect examples, making "
        "unsupervised or one-class anomaly detection appropriate. The MVTec AD dataset provides "
        "a realistic benchmark with pixel-precise annotations for industrial anomalies "
        "(Bergmann et al., 2019; Bergmann et al., 2021). Recent methods such as PatchCore "
        "and PaDiM demonstrate strong performance for detection and localization in this setting "
        "(Roth et al., 2021; Defard et al., 2020)."
    )

    doc.add_heading("2. Aims and Objectives", level=1)
    doc.add_paragraph("Primary aim: build an interpretable, safety-aligned anomaly inspection pipeline.")
    doc.add_paragraph("Objectives:")
    doc.add_paragraph("1) Detect anomalies with image-level decision and continuous score.")
    doc.add_paragraph("2) Localize anomalies with heatmaps, masks, and bounding boxes.")
    doc.add_paragraph("3) Produce defect labels and evidence grounded in visible cues.")
    doc.add_paragraph("4) Compute risk class via deterministic RPM lookup (no VLM risk generation).")
    doc.add_paragraph("5) Provide action recommendations from a fixed policy.")
    doc.add_paragraph("6) Quantify uncertainty and trigger human review when needed.")

    doc.add_heading("3. Research Questions", level=1)
    doc.add_paragraph(
        "RQ1: How effectively do PatchCore and PaDiM localize anomalies on MVTec AD categories?\n"
        "RQ2: Can VLM-based defect labeling provide concise, evidence-grounded semantics without "
        "hallucination, when constrained to fixed label sets?\n"
        "RQ3: How does deterministic risk scoring affect decision transparency and auditability "
        "in industrial inspection pipelines?"
    )

    doc.add_heading("4. Proposed Methodology", level=1)
    doc.add_paragraph("4.1 Dataset")
    doc.add_paragraph(
        "Use the MVTec AD dataset with categories bottle, cable, wood, tile, leather. "
        "Training uses defect-free images only; test images include defects and ground-truth "
        "segmentation masks (Bergmann et al., 2019; Bergmann et al., 2021)."
    )

    doc.add_paragraph("4.2 Anomaly Detection Baselines")
    doc.add_paragraph(
        "Implement PatchCore and PaDiM as strong baselines for detection and localization. "
        "PatchCore uses a memory bank of nominal patch features for outlier scoring, while PaDiM "
        "models patch embeddings with multivariate Gaussians (Roth et al., 2021; Defard et al., 2020)."
    )

    doc.add_paragraph("4.3 Semantics via VLM")
    doc.add_paragraph(
        "Integrate LLaVA for constrained defect labeling. LLaVA is a multimodal model created via "
        "visual instruction tuning, providing strong image-language reasoning (Liu et al., 2023). "
        "The LLaVA-1.6 release improves visual reasoning and OCR, supporting detailed defect evidence "
        "when constrained to fixed label sets (Liu et al., 2024)."
    )

    doc.add_paragraph("4.4 Risk Scoring and Actions")
    doc.add_paragraph(
        "Compute risk class and score using a predefined Risk Priority Matrix (RPM). "
        "The VLM does not generate risk. Action recommendations are mapped deterministically "
        "from risk class."
    )

    doc.add_paragraph("4.5 Uncertainty and Human Review")
    doc.add_paragraph(
        "Combine anomaly model confidence with VLM confidence. Trigger human review when "
        "confidence is low or labels are unknown."
    )

    doc.add_heading("5. System Architecture", level=1)
    doc.add_paragraph(
        "The pipeline is modular: data loader → anomaly model → post-processing → VLM semantics → "
        "risk module → action policy → uncertainty module → reporting. A FastAPI server provides "
        "a user camera UI and admin upload UI."
    )

    doc.add_heading("6. Evaluation Plan", level=1)
    doc.add_paragraph("Image-level: AUROC, AP, F1 at selected thresholds.")
    doc.add_paragraph("Pixel-level: AUROC, PRO, IoU/F1 for masks.")
    doc.add_paragraph("Calibration: ECE or confidence-threshold analysis.")
    doc.add_paragraph(
        "Compare PatchCore and PaDiM on category subsets; analyze sensitivity to thresholding and ROI."
    )

    doc.add_heading("7. Ethical, Safety, and Risk Considerations", level=1)
    doc.add_paragraph(
        "The system avoids hallucinated risk by using deterministic RPM lookup. "
        "Human review is required under uncertainty to reduce unsafe automation. "
        "Data usage follows dataset licensing and attribution requirements."
    )

    doc.add_heading("8. Project Plan (Summary)", level=1)
    doc.add_paragraph(
        "Phase 1: Baseline pipeline and evaluation setup.\n"
        "Phase 2: PatchCore/PaDiM integration and localization metrics.\n"
        "Phase 3: VLM semantics with constrained label sets.\n"
        "Phase 4: Risk scoring, action policy, uncertainty analysis.\n"
        "Phase 5: Final evaluation, reporting, and dissertation writing."
    )

    doc.add_heading("9. Expected Contributions", level=1)
    doc.add_paragraph(
        "1) A reproducible, modular inspection pipeline with interpretable outputs.\n"
        "2) Empirical comparison of anomaly detection baselines on selected MVTec AD categories.\n"
        "3) A safety-aligned risk scoring and action recommendation layer."
    )

    doc.add_heading("References (Harvard Style)", level=1)
    doc.add_paragraph(
        "Bergmann, P., Fauser, M., Sattlegger, D. and Steger, C. (2019) "
        "MVTec AD — A Comprehensive Real-World Dataset for Unsupervised Anomaly Detection. "
        "Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition."
    )
    doc.add_paragraph(
        "Bergmann, P., Batzner, K., Fauser, M., Sattlegger, D. and Steger, C. (2021) "
        "The MVTec Anomaly Detection Dataset: A Comprehensive Real-World Dataset for Unsupervised "
        "Anomaly Detection. International Journal of Computer Vision."
    )
    doc.add_paragraph(
        "Defard, T., Setkov, A., Loesch, A. and Audigier, R. (2020) "
        "PaDiM: a Patch Distribution Modeling Framework for Anomaly Detection and Localization. "
        "arXiv preprint arXiv:2011.08785."
    )
    doc.add_paragraph(
        "Roth, K., Pemula, L., Zepeda, J., Schölkopf, B., Brox, T. and Gehler, P. (2021) "
        "Towards Total Recall in Industrial Anomaly Detection. arXiv preprint arXiv:2106.08265."
    )
    doc.add_paragraph(
        "Liu, H., Li, C., Wu, Q. and Lee, Y.J. (2023) Visual Instruction Tuning. "
        "arXiv preprint arXiv:2304.08485."
    )
    doc.add_paragraph(
        "Liu, H., Li, C., Li, Y., Li, B., Zhang, Y., Shen, S. and Lee, Y.J. (2024) "
        "LLaVA-1.6: Improved reasoning, OCR, and world knowledge (blog release)."
    )

    doc.save(str(out_path))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
